"""
文件处理工具 - 支持多种文档格式的文件处理
"""

import hashlib
import magic
import aiofiles
import tempfile
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
from fastapi import UploadFile, HTTPException

from src.core.logging import get_logger

logger = get_logger(__name__)


class FileProcessingError(Exception):
    """文件处理异常"""
    pass


class FileHandler:
    """文件处理器 - 按照specification规范实现"""

    # 支持的文件类型和MIME类型映射
    SUPPORTED_MIME_TYPES = {
        'text/plain': 'txt',
        'text/markdown': 'md',
        'text/x-markdown': 'md',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/epub+zip': 'epub',
    }

    # 文件扩展名映射
    SUPPORTED_EXTENSIONS = {
        '.txt': 'txt',
        '.md': 'md',
        '.markdown': 'md',
        '.docx': 'docx',
        '.epub': 'epub',
    }

    # 最大文件大小 (50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024

    @classmethod
    def get_file_type_from_extension(cls, filename: str) -> Optional[str]:
        """从文件扩展名获取文件类型 - 按照specification规范实现"""
        if not filename:
            return None

        ext = Path(filename).suffix.lower()
        return cls.SUPPORTED_EXTENSIONS.get(ext)

    @classmethod
    def get_file_type_from_mime(cls, mime_type: str) -> Optional[str]:
        """从MIME类型获取文件类型 - 按照specification规范实现"""
        return cls.SUPPORTED_MIME_TYPES.get(mime_type)

    @classmethod
    async def validate_file(cls, file: UploadFile) -> Tuple[str, Dict[str, Any]]:
        """
        验证上传的文件

        Args:
            file: 上传的文件

        Returns:
            Tuple[文件类型, 文件信息]

        Raises:
            FileProcessingError: 文件验证失败
        """
        if not file.filename:
            raise FileProcessingError("文件名不能为空")

        # 检查文件大小
        file.file.seek(0, 2)  # 移动到文件末尾
        file_size = file.file.tell()
        file.file.seek(0)  # 重置到文件开头

        if file_size > cls.MAX_FILE_SIZE:
            raise FileProcessingError(f"文件大小超过限制，最大允许 {cls.MAX_FILE_SIZE // (1024*1024)}MB")

        if file_size == 0:
            raise FileProcessingError("文件不能为空")

        # 从扩展名推断文件类型
        file_type_ext = cls.get_file_type_from_extension(file.filename)
        if not file_type_ext:
            raise FileProcessingError(f"不支持的文件扩展名: {Path(file.filename).suffix}")

        # 读取文件开头用于MIME类型检测
        file_content = file.file.read(1024)
        file.file.seek(0)  # 重置到文件开头

        # 检测MIME类型
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
            file_type_mime = cls.get_file_type_from_mime(mime_type)
        except Exception as e:
            logger.warning(f"MIME类型检测失败: {e}")
            file_type_mime = None

        # 验证文件类型一致性
        if file_type_mime and file_type_mime != file_type_ext:
            logger.warning(f"文件类型不匹配: 扩展名={file_type_ext}, MIME={file_type_mime or mime_type}")
            # 以MIME类型为准，如果支持的话
            if file_type_mime in cls.SUPPORTED_EXTENSIONS.values():
                file_type = file_type_mime
            else:
                raise FileProcessingError(f"不支持的文件类型: {mime_type}")
        else:
            file_type = file_type_ext

        # 计算文件哈希
        file_hash = await cls.calculate_file_hash(file)
        file.file.seek(0)  # 重置到文件开头

        file_info = {
            'filename': file.filename,
            'size': file_size,
            'content_type': file.content_type,
            'detected_mime': mime_type,
            'file_type': file_type,
            'file_hash': file_hash,
        }

        return file_type, file_info

    @classmethod
    async def calculate_file_hash(cls, file: UploadFile) -> str:
        """计算文件的SHA-256哈希值"""
        file.file.seek(0)
        hash_sha256 = hashlib.sha256()

        # 分块读取文件以处理大文件
        chunk_size = 8192
        while True:
            chunk = file.file.read(chunk_size)
            if not chunk:
                break
            hash_sha256.update(chunk)

        file.file.seek(0)  # 重置到文件开头
        return hash_sha256.hexdigest()

    @classmethod
    async def save_temp_file(cls, file: UploadFile) -> Tuple[str, str]:
        """
        保存文件到临时目录

        Returns:
            Tuple[临时文件路径, 文件名]
        """
        # 创建临时文件
        suffix = Path(file.filename).suffix if file.filename else '.tmp'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            file_path = temp_file.name

        # 保存文件内容
        async with aiofiles.open(file_path, 'wb') as f:
            file.file.seek(0)
            while True:
                chunk = file.file.read(8192)
                if not chunk:
                    break
                await f.write(chunk)

        file.file.seek(0)  # 重置到文件开头
        return file_path, Path(file.filename).name


class TextFileHandler:
    """文本文件处理器"""

    @staticmethod
    async def read_text_file(file_path: str, encoding: str = 'utf-8') -> str:
        """
        读取文本文件

        Args:
            file_path: 文件路径
            encoding: 文件编码

        Returns:
            文件内容
        """
        try:
            async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                content = await f.read()
            return content
        except UnicodeDecodeError:
            # 尝试其他编码
            for alt_encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=alt_encoding) as f:
                        content = await f.read()
                    logger.info(f"使用编码 {alt_encoding} 成功读取文件")
                    return content
                except UnicodeDecodeError:
                    continue

            raise FileProcessingError("无法解码文件内容，尝试了多种编码格式")

    @staticmethod
    def count_words(text: str) -> int:
        """统计字数"""
        # 简单的中文字数统计
        import re
        # 移除HTML标签和特殊字符
        clean_text = re.sub(r'<[^>]+>', '', text)
        clean_text = re.sub(r'[^\w\s\u4e00-\u9fff]', ' ', clean_text)

        # 分别计算中文字符和英文单词
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', clean_text))
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', clean_text))

        return chinese_chars + english_words

    @staticmethod
    def estimate_paragraphs(text: str) -> int:
        """估算段落数量"""
        import re
        # 按照空白行分割段落
        paragraphs = re.split(r'\n\s*\n', text.strip())
        # 过滤空段落
        return len([p for p in paragraphs if p.strip()])


class MarkdownFileHandler:
    """Markdown文件处理器"""

    @staticmethod
    async def read_markdown_file(file_path: str) -> str:
        """读取Markdown文件"""
        return await TextFileHandler.read_text_file(file_path)

    @staticmethod
    def extract_metadata(text: str) -> Dict[str, Any]:
        """提取Markdown元数据"""
        import re

        # 提取标题
        titles = re.findall(r'^#+\s+(.+)$', text, re.MULTILINE)

        # 提取章节标题（# 和 ## 级别）
        chapter_titles = re.findall(r'^#{1,2}\s+(.+)$', text, re.MULTILINE)

        return {
            'titles': titles,
            'chapters': chapter_titles,
            'title_count': len(titles),
            'chapter_count': len(chapter_titles),
        }


class DocxFileHandler:
    """Word文档处理器"""

    @staticmethod
    async def read_docx_file(file_path: str) -> str:
        """读取Word文档"""
        try:
            from docx import Document
            doc = Document(file_path)

            # 提取所有段落文本
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(' | '.join(row_text))

            return '\n\n'.join(text)
        except ImportError:
            raise FileProcessingError("未安装python-docx库，无法处理Word文档")
        except Exception as e:
            raise FileProcessingError(f"读取Word文档失败: {str(e)}")

    @staticmethod
    def extract_structure(file_path: str) -> Dict[str, Any]:
        """提取Word文档结构"""
        try:
            from docx import Document
            doc = Document(file_path)

            headings = []
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    headings.append({
                        'text': paragraph.text,
                        'level': int(level),
                    })

            return {
                'headings': headings,
                'heading_count': len(headings),
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
            }
        except Exception as e:
            logger.error(f"提取Word文档结构失败: {e}")
            return {}


class EpubFileHandler:
    """EPUB文件处理器"""

    @staticmethod
    async def read_epub_file(file_path: str) -> str:
        """读取EPUB文件"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup

            book = epub.read_epub(file_path)
            text_parts = []

            # 提取所有章节内容
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    try:
                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                        # 移除脚本和样式
                        for script in soup(["script", "style"]):
                            script.extract()
                        text = soup.get_text(separator='\n', strip=True)
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"处理EPUB章节失败: {e}")
                        continue

            return '\n\n'.join(text_parts)
        except ImportError:
            raise FileProcessingError("未安装ebooklib和beautifulsoup4库，无法处理EPUB文件")
        except Exception as e:
            raise FileProcessingError(f"读取EPUB文件失败: {str(e)}")

    @staticmethod
    def extract_metadata(file_path: str) -> Dict[str, Any]:
        """提取EPUB元数据"""
        try:
            import ebooklib
            from ebooklib import epub

            book = epub.read_epub(file_path)

            metadata = {}
            # 获取基本元数据
            if book.get_metadata('DC', 'title'):
                metadata['title'] = book.get_metadata('DC', 'title')[0][0]
            if book.get_metadata('DC', 'creator'):
                metadata['creator'] = book.get_metadata('DC', 'creator')[0][0]
            if book.get_metadata('DC', 'language'):
                metadata['language'] = book.get_metadata('DC', 'language')[0][0]

            # 统计章节数
            chapters = [item for item in book.get_items()
                       if item.get_type() == ebooklib.ITEM_DOCUMENT]
            metadata['chapter_count'] = len(chapters)

            return metadata
        except Exception as e:
            logger.error(f"提取EPUB元数据失败: {e}")
            return {}


# 文件处理器工厂
def get_file_handler(file_type: str):
    """获取对应的文件处理器 - 按照specification规范实现"""
    handlers = {
        'txt': TextFileHandler,
        'md': MarkdownFileHandler,
        'docx': DocxFileHandler,
        'epub': EpubFileHandler,
    }

    handler = handlers.get(file_type)
    if not handler:
        raise FileProcessingError(f"不支持的文件类型: {file_type}")

    return handler


__all__ = [
    "FileHandler",
    "TextFileHandler",
    "MarkdownFileHandler",
    "DocxFileHandler",
    "EpubFileHandler",
    "FileProcessingError",
    "get_file_handler",
]